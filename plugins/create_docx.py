"""
plugins/create_docx.py — Generate Word (.docx) documents

Agent calls create_docx with title + sections → gets a formatted .docx file.
No need to know python-docx API — the plugin handles all formatting.

Usage by LLM:
    create_docx({
        "filename": "report.docx",
        "title": "Отчёт по расчёту насыщения ТТ",
        "sections": [
            {"heading": "Введение", "text": "Данный документ описывает..."},
            {"heading": "Методология", "text": "Расчёт выполнен по ГОСТ Р 58669-2019..."},
            {"heading": "Результаты", "text": "Время до насыщения: 45мс..."},
            {"heading": "Выводы", "text": "Трансформатор тока удовлетворяет..."}
        ]
    })
    → "Created report.docx (4 sections, 12.3 KB)"

Requires: pip install python-docx
"""

TOOL_DEFINITION = {
    "name": "create_docx",
    "description": (
        "Create a Word (.docx) document with title and sections. "
        "Pass structured content — the tool handles formatting, fonts, spacing. "
        "Each section has a heading and text. Text can include bullet lists "
        "separated by newlines starting with '- '."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "filename": {
                "type": "string",
                "description": "Output filename, e.g. 'report.docx'",
            },
            "title": {
                "type": "string",
                "description": "Document title (large, centered at the top)",
            },
            "sections": {
                "type": "string",
                "description": (
                    "JSON array of sections. Each: {\"heading\": \"...\", \"text\": \"...\"}. "
                    "Text supports newlines and '- ' for bullet points. "
                    "Example: [{\"heading\":\"Intro\",\"text\":\"Some text...\"},{\"heading\":\"Results\",\"text\":\"- item1\\n- item2\"}]"
                ),
            },
            "author": {
                "type": "string",
                "description": "Author name (optional, goes in document properties)",
            },
        },
        "required": ["filename", "title", "sections"],
    },
}

AGENTS = ["coder", "writer"]


def execute(args: dict) -> str:
    import json
    from pathlib import Path

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return "ERROR: python-docx not installed. Run: pip install python-docx"

    filename = args.get("filename", "document.docx")
    title    = args.get("title", "Untitled")
    author   = args.get("author", "")
    sections_raw = args.get("sections", "[]")

    # Parse sections
    try:
        if isinstance(sections_raw, str):
            sections = json.loads(sections_raw)
        else:
            sections = sections_raw
    except json.JSONDecodeError as e:
        return f"ERROR: Invalid JSON in sections: {e}"

    if not isinstance(sections, list):
        return "ERROR: sections must be a JSON array of {heading, text} objects"

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author if provided
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacer

    # Sections
    for i, sec in enumerate(sections):
        if not isinstance(sec, dict):
            continue

        heading = sec.get("heading", f"Section {i+1}")
        text    = sec.get("text", "")

        doc.add_heading(heading, level=1)

        # Parse text: regular paragraphs and bullet points
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("- ") or line.startswith("• "):
                # Bullet point
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("1.") or line.startswith("2.") or line.startswith("3."):
                # Numbered list
                doc.add_paragraph(line[2:].strip(), style="List Number")
            else:
                doc.add_paragraph(line)

    # Set author in properties
    if author:
        doc.core_properties.author = author

    # Save
    output = Path("workspace") / filename
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

    size = output.stat().st_size
    size_str = f"{size/1024:.1f} KB" if size > 1024 else f"{size} bytes"
    return f"Created {filename} ({len(sections)} sections, {size_str})"
